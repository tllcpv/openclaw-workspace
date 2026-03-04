#!/usr/bin/env python3
"""
资料解析器 - 支持PDF、Excel、Word、图片等多种格式
用于评估项目基础资料的信息提取
"""

import sys
import json
import re
from pathlib import Path
from typing import Dict, List, Optional, Any
import argparse

# 可选依赖，按需导入
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import openpyxl
    from openpyxl import load_workbook
    EXCEL_SUPPORT = True
except ImportError:
    EXCEL_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# OCR 相关依赖
try:
    from pdf2image import convert_from_path
    import pytesseract
    from PIL import Image
    OCR_SUPPORT = True
except ImportError:
    OCR_SUPPORT = False


class DocumentParser:
    """文档解析器主类"""
    
    # 评估资料关键字段定义
    PROPERTY_FIELDS = {
        "basic_info": {
            "property_name": ["项目名称", "房产名称", "标的物名称", "估价对象", "债务企业名称"],
            "property_address": ["项目地址", "房产地址", "地址", "坐落", "地理位置"],
            "property_type": ["房产类型", "房屋类型", "用途", "物业类型", "土地用途"],
            "owner": ["产权人", "所有人", "权利人", "业主", "证载权利人"],
        },
        "physical_info": {
            "building_area": ["建筑面积", "面积", "总面积", "房产面积"],
            "land_area": ["土地面积", "占地面积"],
            "floor_count": ["层数", "总层数", "楼层"],
            "current_floor": ["所在楼层", "层"],
            "build_year": ["建成年份", "竣工年份", "建成年代", "竣工日期", "建成使用日期"],
            "structure": ["建筑结构", "结构类型", "结构"],
        },
        "rights_info": {
            "certificate_no": ["产权证号", "不动产权证号", "房产证号", "证书编号", "权证编号"],
            "land_certificate": ["土地证号", "土地使用权证号"],
            "land_use_type": ["土地用途", "用地性质"],
            "land_use_period": ["土地使用期限", "使用权期限", "终止日期"],
        },
        "rental_info": {
            "current_rent": ["当前租金", "现行租金", "月租金"],
            "lease_start": ["租赁起始日", "起租日期"],
            "lease_end": ["租赁到期日", "到期日期"],
            "tenant": ["承租人", "租户", "租客"],
        }
    }
    
    def __init__(self):
        self.extracted_data = {}
    
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """根据文件类型自动选择解析方法"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            return self.parse_pdf(file_path)
        elif suffix in ['.xlsx', '.xls']:
            return self.parse_excel(file_path)
        elif suffix == '.docx':
            return self.parse_docx(file_path)
        elif suffix in ['.txt', '.csv']:
            return self.parse_text(file_path)
        elif suffix in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp']:
            return self.parse_image(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")
    
    def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """解析PDF文件（支持文本型和扫描件）"""
        if not PDF_SUPPORT:
            return {"error": "未安装pdfplumber，无法解析PDF。请运行: pip install pdfplumber"}
        
        result = {
            "file_type": "pdf",
            "file_path": file_path,
            "pages": 0,
            "full_text": "",
            "extracted_fields": {},
            "tables": [],
            "ocr_used": False
        }
        
        try:
            with pdfplumber.open(file_path) as pdf:
                result["pages"] = len(pdf.pages)
                
                all_text = []
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    all_text.append(text)
                    
                    # 尝试提取表格
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            result["tables"].append({
                                "page": i + 1,
                                "data": table
                            })
                
                result["full_text"] = "\n".join(all_text)
                
                # 如果文字提取为空（扫描件），使用OCR
                if not result["full_text"].strip() and OCR_SUPPORT:
                    result["ocr_used"] = True
                    result["full_text"] = self._ocr_pdf(file_path)
                elif not result["full_text"].strip():
                    result["ocr_error"] = "PDF为扫描件，但未安装OCR依赖。请运行: pip install pytesseract pdf2image pillow"
                
                result["extracted_fields"] = self._extract_fields(result["full_text"])
                
        except Exception as e:
            result["error"] = str(e)
        
        return result
    
    def _ocr_pdf(self, file_path: str) -> str:
        """使用OCR识别PDF中的文字"""
        if not OCR_SUPPORT:
            return ""
        
        try:
            # 将PDF转换为图片
            images = convert_from_path(file_path, dpi=300)
            
            all_text = []
            for i, image in enumerate(images):
                # 使用中文+英文识别
                text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                all_text.append(f"=== 第{i+1}页 ===\n{text}")
            
            return "\n\n".join(all_text)
            
        except Exception as e:
            return f"OCR识别失败: {str(e)}"
    
    def parse_image(self, file_path: str) -> Dict[str, Any]:
        """解析图片文件（JPG/PNG等）"""
        if not OCR_SUPPORT:
            return {"error": "未安装OCR依赖，无法解析图片。请运行: pip install pytesseract pillow"}
        
        result = {
            "file_type": "image",
            "file_path": file_path,
            "ocr_used": True,
            "full_text": "",
            "extracted_fields": {}
        }
        
        try:
            # 打开图片
            image = Image.open(file_path)
            
            # 记录图片信息
            result["image_info"] = {
                "format": image.format,
                "size": image.size,
                "mode": image.mode
            }
            
            # OCR识别
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            result["full_text"] = text
            
            # 提取字段
            result["extracted_fields"] = self._extract_fields(text)
            
        except Exception as e:
            result["error"] = str(e)
        
        return result
    
    def parse_excel(self, file_path: str) -> Dict[str, Any]:
        """解析Excel文件"""
        if not EXCEL_SUPPORT:
            return {"error": "未安装openpyxl，无法解析Excel。请运行: pip install openpyxl"}
        
        result = {
            "file_type": "excel",
            "file_path": file_path,
            "sheets": [],
            "extracted_fields": {}
        }
        
        try:
            wb = load_workbook(file_path, data_only=True)
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_data = {
                    "name": sheet_name,
                    "rows": sheet.max_row,
                    "cols": sheet.max_column,
                    "data": []
                }
                
                # 读取所有单元格数据
                for row in sheet.iter_rows(values_only=True):
                    sheet_data["data"].append(row)
                
                result["sheets"].append(sheet_data)
            
            # 尝试从所有sheet中提取字段
            all_text = self._excel_to_text(result["sheets"])
            result["extracted_fields"] = self._extract_fields(all_text)
            
        except Exception as e:
            result["error"] = str(e)
        
        return result
    
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析Word文件"""
        if not DOCX_SUPPORT:
            return {"error": "未安装python-docx，无法解析Word。请运行: pip install python-docx"}
        
        result = {
            "file_type": "docx",
            "file_path": file_path,
            "paragraphs": [],
            "tables": [],
            "extracted_fields": {}
        }
        
        try:
            doc = Document(file_path)
            
            # 提取段落
            all_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    result["paragraphs"].append(para.text.strip())
                    all_text.append(para.text.strip())
            
            # 提取表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                result["tables"].append(table_data)
            
            result["extracted_fields"] = self._extract_fields("\n".join(all_text))
            
        except Exception as e:
            result["error"] = str(e)
        
        return result
    
    def parse_text(self, file_path: str) -> Dict[str, Any]:
        """解析文本文件"""
        result = {
            "file_type": "text",
            "file_path": file_path,
            "extracted_fields": {}
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            result["content"] = text
            result["extracted_fields"] = self._extract_fields(text)
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(file_path, 'r', encoding='gbk') as f:
                text = f.read()
            result["content"] = text
            result["extracted_fields"] = self._extract_fields(text)
        except Exception as e:
            result["error"] = str(e)
        
        return result
    
    def _extract_fields(self, text: str) -> Dict[str, str]:
        """从文本中提取关键字段"""
        extracted = {}
        
        for category, fields in self.PROPERTY_FIELDS.items():
            extracted[category] = {}
            for field_name, keywords in fields.items():
                value = self._find_field_value(text, keywords)
                if value:
                    extracted[category][field_name] = value
        
        return extracted
    
    def parse_valuation_report(self, file_path: str) -> Dict[str, Any]:
        """解析评估测算报告（一户多房结构）"""
        result = {
            "file_type": "valuation_report",
            "file_path": file_path,
            "projects": []  # 多个项目（多户企业）
        }
        
        # 先按常规方式解析
        base_result = self.parse_file(file_path)
        
        if "error" in base_result:
            return base_result
        
        full_text = base_result.get("full_text", "")
        paragraphs = base_result.get("paragraphs", [])
        tables = base_result.get("tables", [])
        
        # 识别分户标记（如"8-1.", "8-2."等）
        import re
        project_markers = list(re.finditer(r'(\d+-\d+)[\.、]\s*债务企业名称[：:]\s*([^\n]+)', full_text))
        
        if not project_markers:
            # 没有分户标记，按常规单项目处理
            result["projects"].append({
                "project_name": base_result.get("extracted_fields", {}).get("basic_info", {}).get("property_name", ""),
                "properties": [base_result.get("extracted_fields", {})]
            })
            return result
        
        # 解析每个项目（每户企业）
        for i, marker in enumerate(project_markers):
            project_code = marker.group(1)  # 如"8-1"
            company_name = marker.group(2).strip()  # 债务企业名称
            
            # 确定当前项目的文本范围
            start_pos = marker.start()
            end_pos = project_markers[i + 1].start() if i + 1 < len(project_markers) else len(full_text)
            project_text = full_text[start_pos:end_pos]
            
            project_data = {
                "project_code": project_code,
                "project_name": company_name,
                "debtor_company": company_name,
                "properties": []  # 该企业的多套房产
            }
            
            # 从表格中提取房产列表
            for table in tables:
                if isinstance(table, list) and len(table) > 1:
                    headers = table[0] if table else []
                    # 查找包含"资产名称"、"建筑面积"、"权证编号"的表格
                    if any("资产" in str(h) for h in headers) and any("面积" in str(h) for h in headers):
                        for row in table[1:]:
                            if len(row) >= 4:
                                property_data = self._parse_property_from_table_row(row, headers)
                                if property_data.get("property_name"):
                                    property_data["project_name"] = company_name
                                    project_data["properties"].append(property_data)
            
            # 如果没有从表格提取到，尝试从文本提取
            if not project_data["properties"]:
                property_data = self._extract_fields(project_text)
                if property_data.get("basic_info", {}).get("property_name"):
                    property_data["project_name"] = company_name
                    project_data["properties"].append(property_data)
            
            result["projects"].append(project_data)
        
        return result
    
    def _parse_property_from_table_row(self, row: List, headers: List) -> Dict[str, Any]:
        """从表格行解析房产信息"""
        data = {
            "basic_info": {},
            "physical_info": {},
            "rights_info": {}
        }
        
        for i, header in enumerate(headers):
            if i >= len(row):
                break
            
            header_str = str(header).strip()
            value = str(row[i]).strip() if row[i] else ""
            
            # 映射表头到字段
            if "资产名称" in header_str or "抵押物名称" in header_str:
                data["basic_info"]["property_name"] = value
            elif "建筑面积" in header_str or "房产面积" in header_str:
                data["physical_info"]["building_area"] = value
            elif "土地面积" in header_str:
                data["physical_info"]["land_area"] = value
            elif "权证编号" in header_str or "产权证号" in header_str:
                data["rights_info"]["certificate_no"] = value
            elif "权利人" in header_str:
                data["basic_info"]["owner"] = value
        
        return data
    
    def _find_field_value(self, text: str, keywords: List[str]) -> Optional[str]:
        """根据关键词查找字段值"""
        lines = text.split('\n')
        
        for keyword in keywords:
            # 模式1: 关键词: 值
            pattern1 = rf'{keyword}[：:]\s*([^\n]+)'
            match = re.search(pattern1, text)
            if match:
                return match.group(1).strip()
            
            # 模式2: 关键词 值（空格分隔）
            for line in lines:
                if keyword in line:
                    # 尝试从行中提取值
                    parts = line.split(keyword)
                    if len(parts) > 1:
                        value = parts[-1].strip()
                        # 清理常见分隔符
                        value = re.sub(r'^[：:、\s]+', '', value)
                        if value and len(value) < 100:  # 过滤过长内容
                            return value
        
        return None
    
    def _excel_to_text(self, sheets: List[Dict]) -> str:
        """将Excel数据转换为文本以便提取字段"""
        texts = []
        for sheet in sheets:
            for row in sheet.get("data", []):
                row_text = " ".join([str(cell) for cell in row if cell is not None])
                texts.append(row_text)
        return "\n".join(texts)
    
    def get_missing_fields(self, extracted: Dict) -> List[str]:
        """获取缺失的关键字段列表"""
        missing = []
        
        # 定义必填字段
        required_fields = [
            ("basic_info", "property_name", "项目名称"),
            ("basic_info", "property_address", "项目地址"),
            ("physical_info", "building_area", "建筑面积"),
            ("rights_info", "certificate_no", "产权证号"),
        ]
        
        for category, field, label in required_fields:
            if category not in extracted or field not in extracted.get(category, {}):
                missing.append(label)
        
        return missing


def main():
    parser = argparse.ArgumentParser(description='评估资料解析工具')
    parser.add_argument('file', help='要解析的文件路径')
    parser.add_argument('-o', '--output', help='输出JSON文件路径')
    parser.add_argument('--check-missing', action='store_true', help='检查缺失字段')
    parser.add_argument('--report', action='store_true', help='评估报告模式（解析一户多房结构）')
    
    args = parser.parse_args()
    
    doc_parser = DocumentParser()
    
    # 评估报告模式
    if args.report:
        result = doc_parser.parse_valuation_report(args.file)
        
        # 为每个项目检查缺失字段
        for project in result.get("projects", []):
            for prop in project.get("properties", []):
                missing = doc_parser.get_missing_fields(prop)
                prop["missing_fields"] = missing
                required = ["项目名称", "项目地址", "建筑面积", "产权证号"]
                prop["completeness"] = f"{len(required) - len(missing)}/{len(required)}"
    else:
        result = doc_parser.parse_file(args.file)
        
        if args.check_missing and "extracted_fields" in result:
            missing = doc_parser.get_missing_fields(result["extracted_fields"])
            result["missing_fields"] = missing
            required = ["项目名称", "项目地址", "建筑面积", "产权证号"]
            result["completeness"] = f"{len(required) - len(missing)}/{len(required)}"
    
    # 输出JSON
    output_json = json.dumps(result, ensure_ascii=False, indent=2)
    
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(output_json)
        print(f"结果已保存至: {args.output}")
    else:
        print(output_json)


if __name__ == "__main__":
    main()
